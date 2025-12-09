import os
import io
import json
import re
import hashlib
import datetime
import tempfile
import smtplib
from email.mime.text import MIMEText
from collections import Counter

import requests
from tqdm import tqdm
from docx import Document
from pydub import AudioSegment
import whisper

from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload

# ==========================================================
# ⚙️ CONFIGURACIÓN FIJA (TUYA)
# ==========================================================
DRIVE_FOLDER_ID      = "1NUq38acTjxIuEvhPgUvFnFlGEOdw8X5i"     # Carpeta DGAGRO_escucha
SUPABASE_BUCKET      = "dgagro360-transcripciones"
SUPABASE_FOLDER_DOCS = "docx"                                  # carpeta dentro del bucket

FECHA_HOY = datetime.date.today().isoformat()

# ==========================================================
# 🔐 CONFIG DESDE VARIABLES DE ENTORNO (SECRETS GH)
# ==========================================================
SUPABASE_URL = os.environ.get("SUPABASE_URL", "").strip().rstrip("/")
SUPABASE_KEY = os.environ.get("SUPABASE_KEY", "").strip()

GOOGLE_SA_JSON = os.environ.get("GOOGLE_SERVICE_ACCOUNT_JSON", "").strip()

EMAIL_SENDER    = os.environ.get("EMAIL_SENDER", "").strip()
EMAIL_PASSWORD  = os.environ.get("EMAIL_PASSWORD", "").strip()
EMAIL_RECIPIENT = os.environ.get("EMAIL_RECIPIENT", "").strip()

if not SUPABASE_URL or not SUPABASE_KEY:
    raise RuntimeError("❌ Faltan SUPABASE_URL o SUPABASE_KEY en los secrets de GitHub.")

if not GOOGLE_SA_JSON:
    raise RuntimeError("❌ Falta GOOGLE_SERVICE_ACCOUNT_JSON en los secrets de GitHub.")

# ==========================================================
# 🔐 NORMALIZACIÓN AUTOMÁTICA DEL SERVICE ACCOUNT
# ==========================================================
def load_and_fix_service_account_json():
    """
    Repara automáticamente problemas de formato del JSON del Service Account
    cuando viene desde GitHub Secrets:
    - Convierte "\\n" en saltos reales
    - Quita espacios basura
    - Repara el cierre END PRIVATE KEY si está mal pegado
    """
    raw = os.environ.get("GOOGLE_SERVICE_ACCOUNT_JSON", "").strip()
    if not raw:
        raise RuntimeError("❌ Falta GOOGLE_SERVICE_ACCOUNT_JSON en los secrets de GitHub.")

    try:
        info = json.loads(raw)
    except Exception as e:
        raise RuntimeError(f"❌ Error al interpretar GOOGLE_SERVICE_ACCOUNT_JSON: {e}")

    if "private_key" not in info:
        raise RuntimeError("❌ El JSON no contiene 'private_key'.")

    key = info["private_key"]

    # Si tiene \n literales → convertir a saltos reales
    if "\\n" in key:
        key = key.replace("\\n", "\n")

    # Quitar espacios extras
    key = key.strip()

    # Reparar si falta el final
    if "END PRIVATE KEY" not in key:
        # Intento de reparación automática
        key = key.split("-----END PRIVATE KEY-----")[0]
        key += "\n-----END PRIVATE KEY-----\n"

    info["private_key"] = key
    return info


# ==========================================================
# 📂 CLIENTE GOOGLE DRIVE
# ==========================================================
def build_drive_service():
    info = load_and_fix_service_account_json()   # ← USAMOS LA VERSIÓN REPARADA

    scopes = ["https://www.googleapis.com/auth/drive.readonly"]

    try:
        creds = service_account.Credentials.from_service_account_info(
            info,
            scopes=scopes
        )
    except Exception as e:
        print("❌ ERROR DESERIALIZANDO PRIVATE KEY")
        print("Texto recibido:")
        print(info["private_key"])
        raise RuntimeError(f"No se pudo cargar la clave privada: {e}")

    service = build("drive", "v3", credentials=creds, cache_discovery=False)
    return service


# ==========================================================
# ☁️ FUNCIONES SUPABASE STORAGE
# ==========================================================
def supabase_object_url(path: str) -> str:
    """
    path = 'bucket/ruta/del/objeto.ext'
    """
    return f"{SUPABASE_URL}/storage/v1/object/{path}"

def supabase_download(path: str):
    """
    Descarga un objeto de Supabase Storage.
    path ejemplo: 'dgagro360-transcripciones/logs/.processed_log.json'
    Devuelve bytes o None si 404.
    """
    url = supabase_object_url(path)
    headers = {
        "Authorization": f"Bearer {SUPABASE_KEY}",
        "apikey": SUPABASE_KEY,
    }
    r = requests.get(url, headers=headers)
    if r.status_code == 200:
        return r.content
    elif r.status_code == 404:
        return None
    else:
        raise RuntimeError(f"❌ Error al descargar de Supabase [{r.status_code}]: {r.text}")

def supabase_upload(path: str, data: bytes, content_type: str = "application/octet-stream"):
    """
    Sube/actualiza un objeto en Supabase Storage.
    path ejemplo: 'dgagro360-transcripciones/docx/2025-12-08/archivo.docx'
    """
    url = supabase_object_url(path)
    headers = {
        "Authorization": f"Bearer {SUPABASE_KEY}",
        "apikey": SUPABASE_KEY,
        "Content-Type": content_type,
        "x-upsert": "true",
    }
    r = requests.put(url, headers=headers, data=data)
    if r.status_code not in (200, 201, 204):
        raise RuntimeError(f"❌ Error al subir a Supabase [{r.status_code}]: {r.text}")

# ==========================================================
# 🧩 FUNCIONES AUXILIARES LOCALES
# ==========================================================
def cargar_json_remoto_or_default(path: str, default: dict):
    """
    Carga JSON desde Supabase, si no existe devuelve default.
    """
    raw = supabase_download(path)
    if raw is None:
        return default
    try:
        return json.loads(raw.decode("utf-8"))
    except json.JSONDecodeError:
        print(f"⚠️ JSON remoto dañado o vacío en {path}, se reinicia.")
        return default

def guardar_json_remoto(path: str, data: dict):
    supabase_upload(path, json.dumps(data, ensure_ascii=False, indent=2).encode("utf-8"), "application/json")

def duracion_min(path):
    try:
        a = AudioSegment.from_file(path)
        return round(len(a) / 60000, 1)
    except Exception:
        return None

# ==========================================================
# 🧠 DETECCIÓN DE NOMBRE DE CAMPO / LOTE
# ==========================================================
def detectar_nombre_campo(texto, memoria):
    texto_lower = texto.lower()
    patrones = [
        r"campo\s+(?:de\s+)?([A-ZÁÉÍÓÚÑ][a-záéíóúñ]+(?:\s+[A-ZÁÉÍÓÚÑ][a-záéíóúñ]+)*)",
        r"lote\s+(?:de\s+)?([A-ZÁÉÍÓÚÑ][a-záéíóúñ]+(?:\s+[A-ZÁÉÍÓÚÑ][a-záéíóúñ]+)*)",
    ]
    candidatos = []
    for pat in patrones:
        candidatos += re.findall(pat, texto, flags=re.IGNORECASE)

    # Coincidencias con memoria previa de campos
    for conocido in memoria.keys():
        if conocido.lower() in texto_lower:
            candidatos.append(conocido)

    if candidatos:
        nombre = Counter(candidatos).most_common(1)[0][0].strip().title()
        memoria[nombre] = memoria.get(nombre, 0) + 1
        return nombre

    # Fallback: primera palabra capitalizada de 4+ letras
    capitalizadas = re.findall(r"\b[A-ZÁÉÍÓÚÑ][a-záéíóúñ]{3,}\b", texto)
    if capitalizadas:
        posible = capitalizadas[0].title()
        memoria[posible] = memoria.get(posible, 0) + 1
        return posible

    return "Sin_identificar"

# ==========================================================
# 📄 CREAR WORD INDIVIDUAL
# ==========================================================
def crear_docx_audio(salida_path, meta, texto):
    doc = Document()
    doc.add_heading(f"Lomas_Pampeanas – Transcripción: {meta['campo_detectado']}", 0)
    doc.add_paragraph(f"📅 Fecha: {meta['fecha_archivo']}")
    doc.add_paragraph(f"⏱ Duración: {meta['duracion_min']} min")
    doc.add_paragraph(f"📁 Archivo original (Drive): {meta['nombre']}")
    doc.add_paragraph("")
    doc.add_heading("📝 Transcripción completa", level=1)
    doc.add_paragraph(texto)
    doc.save(salida_path)

# ==========================================================
# 📘 CREAR WORD MAESTRO (por día)
# ==========================================================
def crear_docx_maestro(items, salida):
    if os.path.exists(salida):
        doc = Document(salida)
        doc.add_paragraph("")
        doc.add_paragraph("──────────────────────────────")
        doc.add_paragraph(f"Nuevas transcripciones agregadas el {FECHA_HOY}")
        print(f"📎 Actualizando compilado maestro existente: {os.path.basename(salida)}")
    else:
        doc = Document()
        doc.add_heading("Lomas_Pampeanas – Compilado General de Transcripciones", 0)
        doc.add_paragraph(f"Actualizado al {FECHA_HOY}")
        doc.add_paragraph("")
        print(f"📄 Creando nuevo compilado maestro: {os.path.basename(salida)}")

    # Agrupar nuevos ítems por campo detectado
    agrupados = {}
    for it in items:
        agrupados.setdefault(it["campo_detectado"], []).append(it)

    for campo, lista in agrupados.items():
        doc.add_heading(f"📍 {campo}", level=1)
        for it in sorted(lista, key=lambda x: x["fecha_archivo"]):
            doc.add_heading(f"🎧 Audio – {it['fecha_archivo']}", level=2)
            doc.add_paragraph(it["texto"])
            doc.add_paragraph("")

    doc.save(salida)
    print(f"📘 Word maestro actualizado: {salida}")

# ==========================================================
# 📤 ENVÍO DE EMAIL (OPCIONAL)
# ==========================================================
def enviar_email_resumen(subcarpetas_creadas):
    if not (EMAIL_SENDER and EMAIL_PASSWORD and EMAIL_RECIPIENT):
        print("ℹ️ EMAIL_* no configurados; se omite envío de correo.")
        return

    cuerpo_detalles = ""
    for fecha_dia, carpeta_virtual, maestro_path in subcarpetas_creadas:
        maestro_nombre = os.path.basename(maestro_path)
        cuerpo_detalles += (
            f"📅 {fecha_dia}\n"
            f"   ✔ Carpeta virtual Supabase: {carpeta_virtual}\n"
            f"   ✔ Word maestro: {maestro_nombre}\n\n"
        )

    cuerpo = f"""
Hola Damian,

Las transcripciones fueron procesadas correctamente y subidas a Supabase.

Bucket: {SUPABASE_BUCKET}
Carpeta base: {SUPABASE_FOLDER_DOCS}/AAAA-MM-DD/

Detalles generados hoy:

{cuerpo_detalles}

Recordatorio:
- Los DOCX individuales y el maestro por día están organizados por fecha.
- Podes gestionarlos y descargarlos desde la consola de Supabase → Storage.

Saludos,
DG|AGRO360°
"""

    msg = MIMEText(cuerpo, _charset="utf-8")
    msg["Subject"] = "DG|AGRO360° – Nuevas transcripciones procesadas (Supabase)"
    msg["From"] = EMAIL_SENDER
    msg["To"] = EMAIL_RECIPIENT

    try:
        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
            server.login(EMAIL_SENDER, EMAIL_PASSWORD)
            server.sendmail(EMAIL_SENDER, [EMAIL_RECIPIENT], msg.as_string())
        print("📨 Email enviado correctamente.")
    except Exception as e:
        print(f"⚠️ Error al enviar correo: {e}")

# ==========================================================
# 🎙️ LISTAR AUDIOS NUEVOS DESDE DRIVE
# ==========================================================
def listar_audios_drive_nuevos(drive_service, log):
    exts = (".mp3", ".m4a", ".wav", ".ogg", ".flac", ".aac")

    q = f"'{DRIVE_FOLDER_ID}' in parents and trashed = false"
    files = []
    page_token = None
    while True:
        resp = drive_service.files().list(
            q=q,
            spaces="drive",
            fields="nextPageToken, files(id, name, mimeType, modifiedTime)",
            pageToken=page_token,
        ).execute()
        files.extend(resp.get("files", []))
        page_token = resp.get("nextPageToken")
        if not page_token:
            break

    nuevos = []
    for f in files:
        nombre = f["name"]
        if not nombre.lower().endswith(exts):
            continue
        # clave estable por fileId + modifiedTime
        key = f"{f['id']}@{f['modifiedTime']}"
        if key not in log["procesados"]:
            nuevos.append((f, key))

    return nuevos

def descargar_audio_temporal(drive_service, file_id, nombre):
    fd, temp_path = tempfile.mkstemp(prefix="dgagro_audio_", suffix=os.path.splitext(nombre)[1])
    os.close(fd)
    request = drive_service.files().get_media(fileId=file_id)
    fh = io.FileIO(temp_path, "wb")
    downloader = MediaIoBaseDownload(fh, request)
    done = False
    while not done:
        status, done = downloader.next_chunk()
    fh.close()
    return temp_path

# ==========================================================
# 🧵 MAIN
# ==========================================================
def main():
    print("🔄 Iniciando transcriptor DG|AGRO360° con Drive + Supabase...")

    drive_service = build_drive_service()

    # --- Log remoto en Supabase ---
    LOG_PATH_REMOTE = f"{SUPABASE_BUCKET}/logs/.processed_log.json"
    log = cargar_json_remoto_or_default(LOG_PATH_REMOTE, {"procesados": {}})
    memoria_campos = log.get("memoria_campos", {})

    # --- Detectar audios nuevos ---
    pendientes = listar_audios_drive_nuevos(drive_service, log)

    if not pendientes:
        print("✅ No hay audios nuevos para procesar.")
        return

    print(f"🔎 Audios nuevos encontrados: {len(pendientes)}")

    modelo = whisper.load_model("small")
    resumen_items = []

    with tempfile.TemporaryDirectory(prefix="dgagro360_out_") as tmpdir:
        for f, key in tqdm(pendientes, desc="🎧 Transcribiendo"):
            file_id = f["id"]
            nombre = f["name"]
            modified = f["modifiedTime"]

            # Fecha: desde el nombre si tiene, si no, modifiedTime
            fecha_nombre = re.search(r"(\d{4}[-_]\d{2}[-_]\d{2})", nombre)
            if fecha_nombre:
                fecha_archivo = fecha_nombre.group(1).replace("_", "-")
            else:
                # modifiedTime formato ISO: 2025-12-08T14:23:11.000Z
                fecha_archivo = modified[:10]

            # Descargar audio temporal
            path_audio = descargar_audio_temporal(drive_service, file_id, nombre)
            dur_min = duracion_min(path_audio)

            out = modelo.transcribe(path_audio, fp16=False)
            texto = out.get("text", "").strip()

            campo_detectado = detectar_nombre_campo(texto, memoria_campos)
            meta = {
                "nombre": nombre,
                "drive_id": file_id,
                "modifiedTime": modified,
                "fecha_archivo": fecha_archivo,
                "duracion_min": dur_min,
                "campo_detectado": campo_detectado,
                "texto": texto,
            }

            # Crear DOCX individual en tmpdir
            campo_slug = campo_detectado.replace(" ", "_")
            nombre_docx = f"{campo_slug}_{fecha_archivo}.docx"
            local_docx_path = os.path.join(tmpdir, nombre_docx)
            crear_docx_audio(local_docx_path, meta, texto)

            # Guardar en la lista para agrupación por día
            resumen_items.append(meta)

            # Marcar como procesado en el log
            log["procesados"][key] = {
                "nombre": nombre,
                "drive_id": file_id,
                "fecha": fecha_archivo,
                "campo": campo_detectado,
            }

            # Borrar audio temporal
            try:
                os.remove(path_audio)
            except Exception:
                pass

        # ==================================================
        # 📁 AGRUPAR POR FECHA Y SUBIR DOCX A SUPABASE
        # ==================================================
        # Estructura destino:
        #   dgagro360-transcripciones/docx/AAAA-MM-DD/archivo.docx
        subcarpetas_creadas = []
        agrupados_por_fecha = {}
        for it in resumen_items:
            fecha = it["fecha_archivo"]
            agrupados_por_fecha.setdefault(fecha, []).append(it)

        for fecha_dia, lista_dia in agrupados_por_fecha.items():
            carpeta_virtual = f"{SUPABASE_FOLDER_DOCS}/{fecha_dia}"
            # 1) Subir DOCX individuales
            for it in lista_dia:
                campo_slug = it["campo_detectado"].replace(" ", "_")
                nombre_docx = f"{campo_slug}_{fecha_dia}.docx"
                local_docx = os.path.join(tmpdir, nombre_docx)
                if not os.path.exists(local_docx):
                    # Si no existe (caso raro, pero por las dudas)
                    crear_docx_audio(local_docx, it, it["texto"])
                with open(local_docx, "rb") as fbin:
                    remoto_path = f"{SUPABASE_BUCKET}/{carpeta_virtual}/{nombre_docx}"
                    supabase_upload(remoto_path, fbin.read(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

            # 2) Crear y subir Word maestro por fecha
            maestro_local = os.path.join(tmpdir, f"Transcripciones_Completas_{fecha_dia}.docx")
            crear_docx_maestro(lista_dia, maestro_local)
            with open(maestro_local, "rb") as fbin:
                maestro_remoto = f"{SUPABASE_BUCKET}/{carpeta_virtual}/Transcripciones_Completas_{fecha_dia}.docx"
                supabase_upload(maestro_remoto, fbin.read(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

            subcarpetas_creadas.append(
                (fecha_dia, carpeta_virtual, maestro_local)
            )

    # Guardar memoria de campos en el log remoto
    log["memoria_campos"] = memoria_campos
    guardar_json_remoto(LOG_PATH_REMOTE, log)

    print("\n✅ Transcripciones subidas a Supabase correctamente.")
    print(f"📦 Bucket: {SUPABASE_BUCKET} | Carpeta base: {SUPABASE_FOLDER_DOCS}/AAAA-MM-DD/")

    # Enviar email si está configurado
    if subcarpetas_creadas:
        enviar_email_resumen(subcarpetas_creadas)


if __name__ == "__main__":
    main()







